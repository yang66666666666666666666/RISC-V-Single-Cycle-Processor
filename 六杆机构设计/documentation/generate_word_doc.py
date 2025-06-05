#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成Word格式的设计说明书
"""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import pandas as pd

def create_word_document():
    """创建Word文档"""
    doc = Document()
    
    # 设置文档标题
    title = doc.add_heading('六杆机构运动分析与动态静力分析设计说明书', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 1. 设计任务
    doc.add_heading('1. 设计任务', level=1)
    
    doc.add_heading('1.1 设计目的', level=2)
    doc.add_paragraph('对给定的六杆机构进行运动分析和动态静力分析，确定机构的运动规律和受力特性。')
    
    doc.add_heading('1.2 设计要求', level=2)
    requirements = [
        '根据给定参数设计六杆机构的运动尺寸',
        '确定曲柄的转动方向',
        '进行机构运动分析，输出滑块6的运动规律',
        '进行动态静力分析，计算固定铰链反力和平衡力矩',
        '绘制相应的运动曲线和力变化曲线'
    ]
    for req in requirements:
        doc.add_paragraph(req, style='List Bullet')
    
    # 2. 设计参数
    doc.add_heading('2. 设计参数（第9列数据）', level=1)
    
    # 创建参数表格
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '参数名称'
    hdr_cells[1].text = '符号'
    hdr_cells[2].text = '数值'
    hdr_cells[3].text = '单位'
    
    # 参数数据
    parameters = [
        ('转速', 'n₂', '49', 'r/min'),
        ('曲柄长度', 'l₀₂ₐ', '82', 'mm'),
        ('分布', 'd', '122', 'mm'),
        ('行程速比系数', 'K', '1.9', '-'),
        ('描刀行程', 'H', '128', 'mm'),
        ('长度比', 'f_BC/f_BO1', '0.62', '-'),
        ('工作阻力', 'F_max', '8800', 'N'),
        ('滑块6质量', 'm₆', '65', 'kg'),
        ('导杆4质量', 'm₄', '26', 'kg'),
        ('导杆4转动惯量', 'J_s4', '1.2', 'kg·m²')
    ]
    
    for param in parameters:
        row_cells = table.add_row().cells
        for i, value in enumerate(param):
            row_cells[i].text = value
    
    # 3. 设计计算过程
    doc.add_heading('3. 设计计算过程', level=1)
    
    doc.add_heading('3.1 机构尺寸设计', level=2)
    
    doc.add_heading('3.1.1 摆角计算', level=3)
    doc.add_paragraph('根据行程速比系数K，计算导杆的摆角：')
    doc.add_paragraph('K = (180° + ψ) / (180° - ψ)')
    doc.add_paragraph('解得：')
    doc.add_paragraph('ψ = 180° × (K - 1) / (K + 1) = 180° × (1.9 - 1) / (1.9 + 1) = 55.86°')
    
    doc.add_heading('3.1.2 连杆长度计算', level=3)
    doc.add_paragraph('根据行程H和摆角ψ，计算连杆AB长度：')
    doc.add_paragraph('l_AB = H / (2 × sin(ψ/2)) = 128 / (2 × sin(55.86°/2)) = 136.63 mm')
    
    doc.add_heading('3.1.3 其他杆长计算', level=3)
    dimensions = [
        '机架长度：l_O1O2 = d = 122 mm',
        '导杆O1B长度：通过几何约束计算得 l_O1B = 101.88 mm',
        'BC杆长度：l_BC = 0.62 × l_O1B = 63.17 mm',
        'CD杆长度：l_CD = 50.53 mm（经验估算）'
    ]
    for dim in dimensions:
        doc.add_paragraph(dim, style='List Bullet')
    
    doc.add_heading('3.2 转动方向确定', level=2)
    doc.add_paragraph('由于K = 1.9 > 1，说明工作行程时间短于回程时间，为提高工作效率，选择顺时针转动方向。')
    
    # 4. 程序设计
    doc.add_heading('4. 程序设计', level=1)
    
    doc.add_heading('4.1 程序结构', level=2)
    structure = """六杆机构设计/
├── calculations/
│   └── design_calculations.py    # 设计计算模块
├── programs/
│   ├── kinematic_analysis.py     # 运动分析主程序
│   └── dynamic_analysis.py       # 动态静力分析主程序
├── results/                      # 输出结果
└── documentation/               # 说明书"""
    
    doc.add_paragraph(structure, style='Intense Quote')
    
    doc.add_heading('4.2 主要函数说明', level=2)
    
    functions = [
        ('设计计算模块', [
            'SixBarMechanismDesign: 主设计类',
            'calculate_mechanism_dimensions(): 计算机构尺寸',
            'determine_rotation_direction(): 确定转动方向'
        ]),
        ('运动分析模块', [
            'KinematicAnalysis: 运动分析类',
            'position_analysis(): 位置分析',
            'velocity_analysis(): 速度分析',
            'acceleration_analysis(): 加速度分析'
        ]),
        ('动态分析模块', [
            'DynamicAnalysis: 动态分析类',
            'work_resistance_function(): 工作阻力子程序',
            'calculate_inertia_forces(): 惯性力计算',
            'solve_dynamic_equilibrium(): 动态平衡求解'
        ])
    ]
    
    for module_name, func_list in functions:
        doc.add_heading(f'4.2.{functions.index((module_name, func_list))+1} {module_name}', level=3)
        for func in func_list:
            doc.add_paragraph(func, style='List Bullet')
    
    # 5. 计算结果
    doc.add_heading('5. 计算结果', level=1)
    
    doc.add_heading('5.1 机构尺寸参数', level=2)
    
    # 尺寸结果表格
    dim_table = doc.add_table(rows=1, cols=3)
    dim_table.style = 'Table Grid'
    dim_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    dim_hdr = dim_table.rows[0].cells
    dim_hdr[0].text = '参数'
    dim_hdr[1].text = '数值'
    dim_hdr[2].text = '单位'
    
    dim_results = [
        ('连杆AB长度', '136.63', 'mm'),
        ('机架O1O2长度', '122.00', 'mm'),
        ('导杆O1B长度', '101.88', 'mm'),
        ('BC杆长度', '63.17', 'mm'),
        ('CD杆长度', '50.53', 'mm'),
        ('摆角', '55.86', '°')
    ]
    
    for result in dim_results:
        row_cells = dim_table.add_row().cells
        for i, value in enumerate(result):
            row_cells[i].text = value
    
    doc.add_heading('5.2 运动分析结果', level=2)
    motion_results = [
        '滑块6最大位移：约 82 mm',
        '滑块6最大速度：约 2284 mm/s',
        '滑块6最大加速度：约 51301 mm/s²'
    ]
    for result in motion_results:
        doc.add_paragraph(result, style='List Bullet')
    
    doc.add_heading('5.3 动态分析结果', level=2)
    dynamic_results = [
        'O1处最大反力：约 5299 N',
        'O2处最大反力：约 3525 N',
        '最大平衡力矩：约 722 N·m'
    ]
    for result in dynamic_results:
        doc.add_paragraph(result, style='List Bullet')
    
    # 6. 结果分析
    doc.add_heading('6. 结果分析', level=1)
    
    doc.add_heading('6.1 运动特性分析', level=2)
    motion_analysis = [
        '滑块6的位移曲线呈现典型的往复运动特征',
        '速度曲线在工作行程和回程中表现出不同的特性，符合K>1的设计要求',
        '加速度变化较大，需要考虑动态载荷的影响'
    ]
    for analysis in motion_analysis:
        doc.add_paragraph(analysis, style='List Number')
    
    doc.add_heading('6.2 动力特性分析', level=2)
    force_analysis = [
        '固定铰链反力随曲柄转角周期性变化',
        '工作行程中反力较大，主要由工作阻力引起',
        '平衡力矩的变化反映了机构的动态特性'
    ]
    for analysis in force_analysis:
        doc.add_paragraph(analysis, style='List Number')
    
    doc.add_heading('6.3 设计评价', level=2)
    evaluation = [
        '机构设计满足给定的运动要求',
        '动态载荷在可接受范围内',
        '建议在实际应用中考虑减振措施'
    ]
    for eval_item in evaluation:
        doc.add_paragraph(eval_item, style='List Number')
    
    # 7. 结论
    doc.add_heading('7. 结论', level=1)
    conclusions = [
        '成功完成了六杆机构的尺寸设计，确定了各杆件长度',
        '通过运动分析获得了滑块6的完整运动规律',
        '动态静力分析揭示了机构的受力特性',
        '程序计算结果为机构的优化设计提供了理论依据'
    ]
    for conclusion in conclusions:
        doc.add_paragraph(conclusion, style='List Number')
    
    # 8. 参考文献
    doc.add_heading('8. 参考文献', level=1)
    references = [
        '机械原理教程',
        '机构学与机器动力学',
        'Python科学计算与工程应用'
    ]
    for ref in references:
        doc.add_paragraph(ref, style='List Number')
    
    # 添加页脚信息
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('设计者：[姓名]\n完成日期：2024年12月\n版本：V1.0')
    footer_run.font.size = Pt(10)
    
    return doc

def main():
    """主函数"""
    print("正在生成Word文档...")
    
    # 创建文档
    doc = create_word_document()
    
    # 保存文档
    output_dir = os.path.join(os.path.dirname(__file__))
    doc_path = os.path.join(output_dir, '六杆机构设计说明书.docx')
    doc.save(doc_path)
    
    print(f"Word文档已生成：{doc_path}")

if __name__ == "__main__":
    main()